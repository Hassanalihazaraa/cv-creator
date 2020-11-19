from docx import Document
from docx.shared import Inches
import pyttsx3

voice_engine = pyttsx3.init()
new_volume = 0.1
new_voice_rate = 180


def speak(text):
    voice_engine.setProperty('volume', new_volume)
    voice_engine.setProperty('rate', new_voice_rate)
    pyttsx3.speak(text)


document = Document()

# Add profile picture
document.add_picture('img.jpg', width=Inches(2.0))

# inputs from the user
speak('Hi, May i ask your name?')
name = input('What is your name? ')
speak('Hi' + name + ' nice to meet you, What is your phone number?')
phone_number = input(' What is your phone number? ')
speak('What is your email address?')
email = input('What is your email address? ')

# adding all in paragraph
document.add_paragraph(
    name + '\n' + phone_number + '\n' + email
)

# about me
document.add_heading('About me')
speak("Tell me about yourself? ")
document.add_paragraph(
    input("Tell me about yourself? ")
)


# work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

speak('Please enter your work experience')
company = input('Enter a company name ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

speak('Describe your experience at ' + company + ' ')
experience_details = input('Describe your experience at ' + company + ' ')


# more experiences
while True:
    has_more_experiences = input('Do you have more experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter a company name ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ' ')
    else:
        break


# Skills
document.add_heading('Skills')
skill = input('Enter a skill ')
speak('Please enter all your skills')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes':
        skill = input('Enter a skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break


# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Hassan"s app'

# save inside the document
document.save("cv.docx")
